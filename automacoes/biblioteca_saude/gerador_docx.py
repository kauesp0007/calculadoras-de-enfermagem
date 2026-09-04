"""FASE 11 — Geração DOCX (spec §26-31, §54, §56).

Gera um documento Word (.docx) editável por item AUDITED, com estrutura
proporcional ao tipo de material (§27) e proveniência explícita (§46, §63).

- Formato editável (não PDF) — imagens/tabelas podem ser adicionadas depois (§57).
- Metadados do documento Word preenchidos (§56).
- Referências em Vancouver quando presentes.
"""

import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

from .config import JSON_DIR, SAIDA_DOCX_DIR
from .ingestao import carregar_indice

ITENS_DIR = JSON_DIR / "catalogo" / "itens"

NAO_IDENTIFICADO = "não identificado na fonte"


def _slug(titulo: str) -> str:
    s = re.sub(r"[^a-z0-9]+", "-", (titulo or "").lower())
    return s.strip("-")[:60] or "documento"


def _adicionar_lista(doc, titulo: str, itens: list) -> None:
    if not itens:
        return
    doc.add_heading(titulo, level=2)
    for x in itens:
        doc.add_paragraph(str(x), style="List Bullet")


def _gerar_docx(item: dict):
    from docx import Document

    doc = Document()
    titulo = item.get("titulo") or "Documento de conhecimento em saúde"
    doc.add_heading(titulo, level=1)

    # ── Proveniência (fonte original — nunca ocultar §46) ──────────
    autoria = item.get("autoria") or {}
    publicacao = item.get("publicacao") or {}
    doc.add_heading("Fonte original", level=2)
    doc.add_paragraph(f"Instituição: {autoria.get('instituicao') or NAO_IDENTIFICADO}")
    doc.add_paragraph(f"Ano de publicação: {publicacao.get('ano_publicacao') or NAO_IDENTIFICADO}")
    doc.add_paragraph(f"Tipo documental: {item.get('tipo_documental') or NAO_IDENTIFICADO}")
    doc.add_paragraph(f"Identificador da fonte (ID interno): {item.get('fonte_id') or NAO_IDENTIFICADO}")

    if item.get("resumo"):
        doc.add_heading("Resumo", level=2)
        doc.add_paragraph(item["resumo"])

    conteudo = item.get("conteudo") or {}
    _adicionar_lista(doc, "Conceitos fundamentais", conteudo.get("conceitos"))
    _adicionar_lista(doc, "Procedimentos", conteudo.get("procedimentos"))
    _adicionar_lista(doc, "Medicamentos", conteudo.get("medicamentos"))
    _adicionar_lista(doc, "Equipamentos", conteudo.get("equipamentos"))
    _adicionar_lista(doc, "Escalas", conteudo.get("escalas"))
    _adicionar_lista(doc, "Diagnósticos", conteudo.get("diagnosticos"))
    _adicionar_lista(doc, "Intervenções", conteudo.get("intervencoes"))

    evidencia = item.get("evidencia") or {}
    for titulo_sec, campo in (
        ("Metodologia", "metodologia"),
        ("Resultados", "resultados"),
        ("Conclusões", "conclusoes"),
        ("Recomendações", "recomendacoes"),
    ):
        if evidencia.get(campo):
            doc.add_heading(titulo_sec, level=2)
            doc.add_paragraph(evidencia[campo])

    referencias = item.get("referencias") or []
    if referencias:
        doc.add_heading("Referências", level=2)
        for r in referencias:
            doc.add_paragraph(r.get("texto_formatado") or r.get("titulo") or "")

    # ── Nota de transparência (§69) ────────────────────────────────
    doc.add_heading("Nota de transparência", level=2)
    doc.add_paragraph(
        "Este documento é uma síntese gerada automaticamente a partir da fonte "
        f"indicada acima. Origem dos metadados: {item.get('metadados_origem') or NAO_IDENTIFICADO}. "
        f"Data de processamento: {item.get('data_processamento') or NAO_IDENTIFICADO}."
    )

    # ── Metadados do documento Word (§56) ──────────────────────────
    classificacao = item.get("classificacao") or {}
    assuntos = classificacao.get("assuntos") or []
    cp = doc.core_properties
    cp.title = titulo
    cp.subject = ", ".join(assuntos)
    cp.author = autoria.get("instituicao") or ""
    cp.keywords = ", ".join(assuntos)

    return doc


def executar(dry_run: bool = False) -> dict:
    resumo = {"gerados": [], "pulados": []}

    itens = []
    if ITENS_DIR.exists():
        for p in ITENS_DIR.glob("*.json"):
            itens.append(json.loads(p.read_text(encoding="utf-8")))

    for item in itens:
        if item.get("status") != "AUDITED":
            resumo["pulados"].append(item.get("titulo") or item["id"])
            continue

        nome_arquivo = f"{_slug(item.get('titulo'))}-{item['id'][-8:]}.docx"
        resumo["gerados"].append(nome_arquivo)

        if not dry_run:
            doc = _gerar_docx(item)
            SAIDA_DOCX_DIR.mkdir(parents=True, exist_ok=True)
            doc.save(str(SAIDA_DOCX_DIR / nome_arquivo))

            # marca item como gerado
            item["docx_gerado_em"] = f"biblioteca_de_enfermagem/{nome_arquivo}"
            item["status"] = "WRITING"
            (ITENS_DIR / f"{item['id']}.json").write_text(
                json.dumps(item, ensure_ascii=False, indent=2), encoding="utf-8"
            )

    return resumo


def _main() -> int:
    dry_run = "--dry-run" in sys.argv
    r = executar(dry_run=dry_run)
    modo = "DRY-RUN" if dry_run else "EXECUÇÃO"
    print(f"[FASE 11 — DOCX] modo={modo}")
    print(f"  Gerados : {len(r['gerados'])}")
    for n in r["gerados"]:
        print(f"      - {n}")
    print(f"  Pulados : {len(r['pulados'])}")
    return 0


if __name__ == "__main__":
    raise SystemExit(_main())
